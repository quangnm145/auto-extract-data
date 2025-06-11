import argparse
import msoffcrypto
from docx import Document
import io
import sys
import os

# Thiết lập phân tích tham số dòng lệnh
parser = argparse.ArgumentParser(
    description='Chuyển file .docx (có hoặc không có mật khẩu) sang .md, bao gồm cả bảng',
    epilog='Ví dụ: python docx_to_md.py -i test.docx -c 1a@ -o test.md\n'
           '     python docx_to_md.py -i test.docx\n'
           '     python docx_to_md.py -h'
)
parser.add_argument('-i', '--input', help='File .docx đầu vào')
parser.add_argument('-c', '--password', help='Mật khẩu của file .docx (nếu có)')
parser.add_argument('-o', '--output', help='File .md đầu ra (mặc định: tên file đầu vào với đuôi .md)')

# Kiểm tra nếu không có tham số hoặc chỉ có -h
if len(sys.argv) == 1:
    parser.print_help()
    sys.exit(0)

args = parser.parse_args()

# Kiểm tra nếu thiếu file đầu vào
if not args.input:
    print("Lỗi: Cần cung cấp file đầu vào (-i).")
    parser.print_help()
    sys.exit(1)

# Xác định file đầu ra
if args.output:
    output_file = args.output
else:
    output_file = os.path.splitext(args.input)[0] + '.md'

try:
    # Mở file .docx
    with open(args.input, 'rb') as file:
        office_file = msoffcrypto.OfficeFile(file)
        
        if office_file.is_encrypted():
            if not args.password:
                print("Lỗi: File được mã hóa, cần cung cấp mật khẩu (-c).")
                sys.exit(1)
            
            # Nạp mật khẩu và giải mã
            try:
                office_file.load_key(password=args.password)
                decrypted = io.BytesIO()
                office_file.decrypt(decrypted)
                print("Giải mã thành công.")
                document = Document(decrypted)
            except Exception as e:
                print(f"Giải mã thất bại: {e}")
                sys.exit(1)
        else:
            # File không mã hóa, sử dụng trực tiếp
            print("File đã qua xác thực, tiếp tục chuyển đổi.")
            document = Document(args.input)
        
        # Trích xuất nội dung theo thứ tự
        content = []
        para_index = 0
        table_index = 0
        
        # Duyệt qua các phần tử trong body của tài liệu
        for element in document.element.body:
            if element.tag.endswith('p'):  # Đoạn văn
                if para_index < len(document.paragraphs):
                    para = document.paragraphs[para_index]
                    if para.text.strip():  # Bỏ qua đoạn trống
                        content.append(para.text)
                        content.append('')  # Thêm dòng trống cho Markdown
                    para_index += 1
            elif element.tag.endswith('tbl'):  # Bảng
                if table_index < len(document.tables):
                    table = document.tables[table_index]
                    # Trích xuất dữ liệu bảng
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                        table_data.append(row_data)
                    
                    # Tạo bảng Markdown
                    if table_data:
                        # Thêm hàng tiêu đề
                        content.append('|' + '|'.join(table_data[0]) + '|')
                        # Thêm hàng phân tách
                        content.append('|' + '|'.join(['-' * max(3, len(cell)) for cell in table_data[0]]) + '|')
                        # Thêm các hàng dữ liệu còn lại
                        for row in table_data[1:]:
                            content.append('|' + '|'.join(row) + '|')
                        # Thêm dòng trống sau bảng
                        content.append('')
                    table_index += 1
        
        # Nếu còn đoạn văn hoặc bảng chưa xử lý (hiếm), thêm vào
        while para_index < len(document.paragraphs):
            para = document.paragraphs[para_index]
            if para.text.strip():
                content.append(para.text)
                content.append('')
            para_index += 1
        
        while table_index < len(document.tables):
            table = document.tables[table_index]
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                content.append('|' + '|'.join(table_data[0]) + '|')
                content.append('|' + '|'.join(['-' * max(3, len(cell)) for cell in table_data[0]]) + '|')
                for row in table_data[1:]:
                    content.append('|' + '|'.join(row) + '|')
                content.append('')
            table_index += 1
        
        # Nối nội dung, loại bỏ dòng trống thừa ở cuối
        text = '\n'.join(line for line in content if line or content.index(line) < len(content)-1)
        
        # Ghi văn bản vào file .md
        with open(output_file, 'w', encoding='utf-8') as md_file:
            md_file.write(text)
        print(f"Chuyển đổi hoàn tất, lưu tại: {output_file}")

except Exception as e:
    print(f"Không thể xử lý file: {e}")
    sys.exit(1)